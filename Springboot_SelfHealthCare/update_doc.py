"""
毕业设计手册文档更新脚本
根据实际代码内容，将文档中的内容纠正并扩充到与现有代码版本一致。

主要修改点：
1. 任务书中"成员档案管理"→"个人健康档案管理"（项目是单用户个人健康管理，非家庭成员管理）
2. 修正"成员总数"等家庭成员相关描述
3. 补充文档导入功能（DocumentImportService，支持PDF/图片OCR）
4. 补充数据可视化模块（VisualizationController/Service）
5. 扩充各阶段工作进程记录（更详细的技术描述）
6. 修正技术指标描述（添加PDFBox、Tesseract OCR、BCrypt等）
"""

import docx
from docx import Document
from docx.shared import Pt
import copy

doc = Document('基于SpringBoot的个人健康管理与疾病预警系统.docx')


def replace_in_cell(cell, old_text, new_text):
    """替换单元格中的文本，尽量保留格式"""
    for para in cell.paragraphs:
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
    # 如果 runs 中没有找到，尝试整段处理
    full_text = cell.text
    if old_text in full_text and not any(old_text in run.text for para in cell.paragraphs for run in para.runs):
        for para in cell.paragraphs:
            if old_text in para.text:
                # 清除所有run，重写
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = para.text.replace(old_text, new_text)


def set_cell_text(cell, new_text):
    """完全替换单元格文本内容"""
    # 保留第一个段落的格式，清除其他段落
    first_para = cell.paragraphs[0]
    
    # 清空所有段落
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    
    # 删除多余段落（保留第一个）
    while len(cell.paragraphs) > 1:
        para = cell.paragraphs[-1]
        para._element.getparent().remove(para._element)
    
    # 设置第一段内容
    if first_para.runs:
        first_para.runs[0].text = new_text
    else:
        first_para.add_run(new_text)


def get_table_cell(table_idx, row_idx, col_idx):
    return doc.tables[table_idx].rows[row_idx].cells[col_idx]


# ======================================================
# Table 0: 任务书
# ======================================================
table0 = doc.tables[0]

# Row 6 [0]: 课题来源、目的与要求
# 修改"家庭成员" -> "个人用户"，修改功能描述与实际代码一致
new_row6_text = """一、课题来源、目的与要求：

课题来源：解决社会实际问题的课题。

目的：围绕个人日常健康管理需求，设计并实现一个基于 Spring Boot 的个人健康管理与疾病预警系统。系统以个人健康档案、健康记录、风险评估、预警提示、数据可视化和健康档案详情分析为主线，解决健康数据分散、持续记录困难、异常指标不易被及时识别的问题，并通过集成 DeepSeek AI 接口提供智能健康咨询支持。

要求：完成系统的分析、设计、实现与测试。系统应实现用户注册登录、个人健康档案管理、健康记录管理（含20余项生理与生活方式指标）、风险评估与预警、健康总览仪表盘、档案详情与趋势分析、健康数据可视化、健康文档导入（支持PDF/OCR图片解析）以及可选 AI 咨询接口等功能；采用前后端分离式接口组织方式，后端使用 Spring Boot 3.5.11、Spring MVC、Spring Data JPA、Spring Security Crypto，数据库采用 H2 文件数据库（MySQL兼容模式）；前端使用 HTML、CSS、JavaScript 完成页面展示与交互；在毕业设计过程中完成需求分析、总体设计、详细设计、功能测试及设计说明书撰写。"""

for cell in table0.rows[6].cells:
    if cell.text.strip().startswith('一、课题来源'):
        set_cell_text(cell, new_row6_text)
        break

# Row 7 [0]: 主要设计内容
new_row7_text = """二、主要设计内容：

本课题《基于SpringBoot的个人健康管理与疾病预警系统》主要设计内容如下：

1、用户认证模块
实现用户注册、登录、退出及会话状态查询，采用 BCryptPasswordEncoder 对密码进行加密存储，通过 HttpSession 管理登录状态，保障数据隔离与访问安全。

2、健康总览仪表盘模块
展示当前用户的健康记录总数、待处理预警数、高风险记录数、最新风险等级、风险等级分布、最近5条健康记录和最近6条预警信息，为用户提供健康状态全局概览。

3、个人健康档案管理模块
实现用户个人健康档案的创建与更新（upsert语义），维护真实姓名、性别、年龄、出生日期、血型、身高体重、吸烟饮酒状态、家族病史、慢性病史、过敏信息、当前用药、手术史、运动习惯、健康目标、紧急联系人等完整档案信息；同时提供档案详情接口，聚合近期趋势、个性化健康建议与最近预警。

4、健康记录管理模块
实现体重、腰围、血压（收缩压/舒张压）、心率、空腹血糖、餐后血糖、体温、血氧、总胆固醇、BMI（自动计算）、睡眠时长、运动时长、步数、饮水量、压力等级、情绪评分、症状描述、用药记录等20余项指标的录入、修改、删除和分页筛选展示；记录保存后自动触发风险评估与预警刷新。

5、风险评估与预警模块
依据 RiskAssessmentService 中的多指标阈值规则对 BMI、血压、空腹血糖、餐后血糖、血氧、总胆固醇、腰围（区分性别）、睡眠、压力、步数和情绪等指标进行评估，生成风险评分、风险等级（LOW/MEDIUM/HIGH/CRITICAL）、系统摘要和预警草稿；预警信息按严重程度分级（LOW/MEDIUM/HIGH/CRITICAL），支持状态更新与处理备注。

6、健康数据可视化模块
提供最近20条记录的多维度趋势数据，包括体重、收缩压、空腹血糖、睡眠时长、运动时长和风险分数6个指标系列，每系列包含趋势数据点、最新值和平均值，并提供风险等级分布及近10条记录摘要，供前端图表渲染。

7、健康文档导入模块
支持 PDF（基于 Apache PDFBox）、图片（基于 Tesseract OCR，支持中英文识别）和纯文本三类健康文档的上传与自动解析，通过正则表达式提取26项健康字段，将解析结果合并入个人档案并自动创建健康记录。

8、档案详情与 AI 咨询模块
聚合个人档案、最新健康记录、6项近期趋势指标、个性化健康建议和近期预警；在配置 DEEPSEEK_API_KEY 环境变量的情况下，支持向 DeepSeek Chat API 发起 AI 健康分析请求，返回中文分析结果与免责声明。"""

for cell in table0.rows[7].cells:
    if cell.text.strip().startswith('二、主要设计内容'):
        set_cell_text(cell, new_row7_text)
        break

# Row 8 [0]: 主要设计技术指标与参数
new_row8_text = """三、主要设计技术指标与参数：
1、开发语言与平台：Java 21，Spring Boot 3.5.11。
2、后端框架：Spring MVC、Spring Data JPA（Hibernate）、Jakarta Validation、Spring Security Crypto（BCrypt密码加密）。
3、数据库：H2 文件数据库（jdbc:h2:file模式，MySQL兼容模式，AUTO_SERVER=TRUE），支持本地持久化演示；通过 spring.jpa.hibernate.ddl-auto=update 自动管理表结构。
4、前端实现：HTML + CSS + JavaScript，通过 fetch 调用 REST API，支持表单验证、弹窗交互和前端分页。
5、文档处理：Apache PDFBox 3.0.5（PDF文本提取）；Tesseract OCR（系统级，支持 chi_sim+eng 中英文识别）。
6、AI集成：DeepSeek Chat API（通过 Spring RestClient 调用），temperature=0.4，max_tokens=900，支持通过环境变量 DEEPSEEK_ENABLED/DEEPSEEK_API_KEY 控制开关。
7、系统结构：controller、service、repository、domain、dto、exception、util 分层组织，遵循单一职责原则。
8、安全机制：基于 HttpSession 的 Session-Cookie 认证，密码 BCrypt 加密，所有 Service 方法均执行所有权校验（userId匹配）防止越权访问。
9、核心技术要求：实现健康记录保存后自动触发风险评估（RiskAssessmentService）、刷新预警信息（HealthAlertRepository），并支持档案详情聚合分析（ProfileInsightService）与可选 AI 咨询（DeepSeekConsultationService）。
10、完成系统测试和设计说明书撰写，保证主要功能可运行、可演示、可复现。"""

for cell in table0.rows[8].cells:
    if cell.text.strip().startswith('三、主要设计技术指标'):
        set_cell_text(cell, new_row8_text)
        break


# ======================================================
# Table 2: 开题报告
# ======================================================
table2 = doc.tables[2]

# Row 2: 课题的目的意义
new_t2r2_text = """一、课题的目的意义：
目的：基于 Spring Boot 开发一个面向个人用户使用的健康管理与疾病预警系统，实现个人健康档案维护、健康记录持续记录、异常指标自动识别、风险等级评估与预警提示，帮助用户形成持续记录与主动干预的健康管理方式。系统在支持手动录入健康数据的基础上，还提供健康文档导入功能（PDF/图片OCR解析），并集成 DeepSeek AI 接口提供智能问诊支持。
意义：一方面，该系统能够把分散的个人健康信息集中管理，并通过多指标规则化评估及时提示潜在风险；系统对BMI、血压、血糖、血氧、胆固醇等12类生理指标建立了分级阈值判断模型，能够在每次健康记录保存时自动触发评估并刷新预警。另一方面，项目采用 Spring Boot 分层架构、RESTful API、JPA 持久化、Session认证、PDF/OCR文档处理和外部AI服务集成，技术栈较为完整，具有较好的工程训练价值，适合作为软件工程专业毕业设计课题。"""

for cell in table2.rows[2].cells:
    if cell.text.strip().startswith('一、课题的目的意义'):
        set_cell_text(cell, new_t2r2_text)
        break

# Row 3: 资料调研分析
new_t2r3_text = """二、资料调研分析：
（一）现有健康管理类产品与系统调研

目前市面上的健康管理类应用主要分为以下几类：第一类是以微信/支付宝小程序为载体的轻量健康记录工具（如"血压记录本"、"血糖日记"等），功能单一，仅支持单项指标录入，缺乏跨指标综合评估能力，且数据私有性依赖第三方平台；第二类是医院/医疗机构提供的健康档案系统（如电子病历平台、区域健康信息平台），功能完善但面向医疗机构专业人员，普通个人用户入口受限，且通常不开放接口供学习和二次开发；第三类是面向消费者的商业健康APP（如华为运动健康、Apple健康、Keep等），功能丰富但高度依赖硬件设备（智能手环/手表）采集数据，不适合以手动输入为主的学习场景，源码也不对外开放。

（二）后端框架选型调研

在Java Web框架领域，目前主流选择包括 Spring Boot、Quarkus 和 Micronaut。Quarkus 和 Micronaut 更适合云原生/微服务场景，学习曲线较陡；Spring Boot 生态成熟、文档丰富、教学资源充足，是目前高校教学和企业项目中使用最广泛的Java后端框架。Spring Data JPA 在 Spring Boot 体系下提供了"约定优于配置"的数据访问封装，相比直接使用 MyBatis 减少了大量样板SQL代码，适合以对象模型为中心的领域设计。因此本项目选用 Spring Boot 3.5.11 + Spring Data JPA 作为后端核心框架。

（三）数据库选型调研

生产级项目通常选用 MySQL、PostgreSQL 等独立数据库服务。但本项目定位为毕业设计演示系统，为降低部署门槛、简化环境配置，调研了 H2 Database 的适用性：H2 支持文件持久化模式（jdbc:h2:file），兼容 MySQL 方言，可通过 spring.jpa.hibernate.ddl-auto=update 自动维护表结构，且内置 Web 控制台便于数据可视化查看，无需额外安装数据库服务即可一键启动演示，完全满足本项目的教学演示需求。

（四）文档处理与AI集成调研

在健康档案导入场景中，用户可能持有PDF或图片格式的体检报告。调研了以下方案：Apache PDFBox 是Apache基金会维护的开源PDF处理库，无需依赖外部服务即可在本地完成PDF文本提取，License 开放；Tesseract OCR 是Google开源的主流OCR引擎，支持100+语言（含中文简体chi_sim包），可通过系统级命令行调用，免费且离线可用。在AI咨询方面，调研了 OpenAI ChatGPT API、百度文心一言API 和 DeepSeek Chat API：DeepSeek 提供与 OpenAI 兼容的接口规范，接入成本低，中文理解能力出色，且提供较为优惠的API调用价格，适合学生项目集成。综合以上调研，本项目采用 PDFBox + Tesseract + DeepSeek Chat API 的技术组合。

（五）前端技术选型调研

当前主流前端框架包括 Vue.js、React 和 Angular。三者均需要 Node.js 工具链支持，构建配置复杂，不利于快速验证后端接口逻辑。对于以后端为重点的毕业设计项目，采用原生 HTML + CSS + JavaScript（通过 fetch 调用 REST API）能够显著降低工程配置成本，将重心聚焦于后端业务逻辑和系统架构设计，且便于在无构建工具环境下直接运行演示。"""

for cell in table2.rows[3].cells:
    if cell.text.strip().startswith('二、资料调研分析'):
        set_cell_text(cell, new_t2r3_text)
        break

# Row 4: 设计方案可行性分析
new_t2r4_text = """三、设计方案的可行性分析和预期目标：
设计方案：系统以后端 Spring Boot 应用为核心，提供用户认证、个人健康档案、健康记录、预警中心、健康总览仪表盘、数据可视化、文档导入和档案详情等 REST 接口；数据层采用 Spring Data JPA 访问 H2 文件数据库（MySQL兼容模式）；安全层采用 BCryptPasswordEncoder 加密密码，通过 HttpSession 管理登录状态；文档处理集成 Apache PDFBox 和 Tesseract OCR；前端使用 HTML、CSS、JavaScript 构建页面并调用接口；在服务器配置有效 DEEPSEEK_API_KEY 环境变量时，可通过 DeepSeekConsultationService 向 DeepSeek Chat API 发起 AI 健康咨询请求。
可行性分析：
1、技术可行性：项目已具备完整源码（61个Java文件），后端依赖清晰，使用 Spring Boot 3.5.11、Java 21、JPA 和 H2 数据库，文档处理采用成熟的 PDFBox 3.0.5 和 Tesseract OCR，页面交互由原生 JavaScript 实现，部署与演示门槛较低，通过 start-app.bat 脚本即可一键启动。
2、经济可行性：系统可在普通个人电脑上运行（服务端口8081），无需额外商业数据库和前端框架投入，H2 数据库以文件形式本地持久化（data/self-health-care-v2），适合教学与答辩演示场景。
3、操作可行性：系统页面按"总览—档案—记录—预警—可视化—详情"组织，交互流程明确，用户可通过表单完成数据录入与查看；H2 控制台（/h2-console）支持数据库可视化查看。
预期目标：完成用户认证、健康总览、个人档案管理、健康记录管理、风险评估与预警、数据可视化、健康文档导入以及可选 AI 咨询功能，实现主要业务流程联通并完成测试与文档撰写。"""

for cell in table2.rows[4].cells:
    if cell.text.strip().startswith('三、设计方案的可行性'):
        set_cell_text(cell, new_t2r4_text)
        break


# ======================================================
# Table 3: 工作进程记录1（第1周-第3周）
# ======================================================
table3 = doc.tables[3]

new_t3r1_text = """本阶段完成的具体工作回顾及完成情况自我评价。
工作回顾：
本阶段重点围绕课题调研与项目源码全面熟悉展开。通过依次阅读 pom.xml、application.properties、SelfHealthCareApplication.java 以及 controller、service、repository、domain、dto、util 等各目录文件，系统梳理了项目的完整技术栈与功能模块。
主要梳理内容包括：
（1）技术栈确认：后端采用 Spring Boot 3.5.11、Java 21、Spring Data JPA（Hibernate ORM）、Spring Security Crypto（BCrypt密码加密）、Apache PDFBox 3.0.5，数据库使用 H2 文件数据库（MySQL兼容模式，AUTO_SERVER=TRUE），前端采用原生 HTML/CSS/JavaScript。
（2）领域模型分析：梳理了 BaseEntity（公共字段基类，含 id/createdAt/updatedAt）、AppUser（用户账号，含 username/displayName/passwordHash）、UserProfile（个人健康档案，与 AppUser 1:1关联，含24个字段）、HealthRecord（健康记录，与 AppUser 多:1关联，含22个生理与生活方式指标字段）、HealthAlert（健康预警，与 HealthRecord 多:1关联，含 category/title/severity/status 等字段）等5个核心实体及其关联关系。
（3）功能模块梳理：确认系统已实现用户注册/登录（Session-Cookie认证）、个人档案管理、健康记录管理（CRUD+分页+风险筛选）、风险评估与预警生成、健康总览仪表盘、数据可视化、健康文档导入（PDF/OCR/文本）、档案详情聚合分析以及 DeepSeek AI 咨询接口共8大功能模块，对应8个 Controller 类。
（4）API端点整理：汇总了 /api/auth（4个端点）、/api/records（5个端点）、/api/alerts（4个端点）、/api/dashboard（1个端点）、/api/profile（3个端点）、/api/visualization（1个端点）、/api/ai（1个端点）、/api/imports（1个端点）共20个REST API端点。
据此完成了课题的初步需求分析和开题报告撰写。
自我评价：
能够按计划完成资料查阅和代码结构梳理工作，对系统整体技术路线和业务闭环有了较清晰认识。通过对源码逐步阅读，避免了仅凭题目想象系统功能的问题，尤其是通过实际阅读 DocumentImportService 和 VisualizationService，发现系统的文档导入和可视化功能在初步预想中未充分体现，从而在开题报告中予以补充，为后续设计和实现分析打下了坚实基础。"""

for cell in table3.rows[1].cells:
    if cell.text.strip().startswith('本阶段完成的具体工作'):
        set_cell_text(cell, new_t3r1_text)
        break


# ======================================================
# Table 4: 工作进程记录2（第4周-第6周）
# ======================================================
table4 = doc.tables[4]

new_t4r1_text = """本阶段完成的具体工作回顾及完成情况自我评价。
工作回顾：
在上一阶段源码熟悉的基础上，本阶段进一步完成了系统需求分析和总体设计工作，重点围绕数据模型关系、接口层次划分与前后端交互方式展开。
主要工作内容如下：
（1）数据关系分析：深入分析了4张数据库表之间的关系——app_user（用户账号）与 user_profile（个人档案）为1:1关系；app_user 与 health_record（健康记录）为1:N关系；health_record 与 health_alert（健康预警）为1:N关系，即每条记录保存时会根据评估结果刷新对应的预警列表。明确了 BaseEntity 中 id/createdAt/updatedAt 字段由 @PrePersist/@PreUpdate 生命周期方法自动维护。
（2）DTO 层分析：梳理了 AuthLoginRequest/AuthRegisterRequest/AuthSessionResponse、HealthRecordRequest/HealthRecordResponse、HealthAlertResponse/AlertStatusUpdateRequest、UserProfileRequest/UserProfileResponse/ProfileDetailResponse、DashboardSummaryResponse、VisualizationResponse/VisualizationSeriesResponse/TrendPointResponse、TrendMetricResponse、AiConsultationRequest/AiConsultationResponse、ImportResultResponse、PagedResponse 共19个DTO类的字段与用途。
（3）接口层次分析：明确了 Controller → Service → Repository 三层调用关系；分析了 AuthService 如何通过 HttpSession 管理登录状态（SESSION_USER_ID 常量），以及 requireAuthenticatedUser() 方法如何在所有业务方法中强制要求用户已登录；分析了 PagingSupport 工具类如何统一处理分页参数归一化与响应封装。
（4）风险评估总体设计：分析了 RiskAssessmentService.assess() 方法的总体评估流程——依次对BMI、血压、空腹血糖、餐后血糖、血氧、总胆固醇、腰围（区分性别）、睡眠、压力、步数、情绪11个维度进行阈值判断，每个维度生成0-3分的风险分，最终汇总得出总分并映射到 LOW/MEDIUM/HIGH/CRITICAL 四档风险等级，同时生成 AlertDraft 列表传递给 HealthRecordService 用于生成 HealthAlert 记录。
自我评价：
能够将业务需求与代码模块相对应，逐步形成系统总体设计思路，对接口分层、实体关系和页面组织方式理解更加深入。特别是通过分析 RiskAssessmentService 中11个维度20+个阈值常量的设计，理解了规则引擎式评估模型的实现方式，为详细设计阶段撰写评估算法说明奠定了基础。"""

for cell in table4.rows[1].cells:
    if cell.text.strip().startswith('本阶段完成的具体工作'):
        set_cell_text(cell, new_t4r1_text)
        break


# ======================================================
# Table 5: 工作进程记录3（第7周-第9周）
# ======================================================
table5 = doc.tables[5]

new_t5r1_text = """本阶段完成的具体工作回顾及完成情况自我评价。
工作回顾：
本阶段主要完成系统详细设计分析与核心代码实现内容的整理工作，深入研读了各Service类和关键工具类的实现细节。
主要工作内容如下：
（1）HealthRecordService 详细分析：分析了 saveRecord() 私有方法中的完整业务链路——将 HealthRecordRequest 中的字段通过 apply() 方法映射到 HealthRecord 实体，调用 RiskAssessmentService.assess() 获取评估结果后设置 bmi/riskScore/riskLevel/summary 字段，保存记录后调用 refreshAlerts() 方法先删除该记录的全部旧预警、再批量保存新生成的 HealthAlert 列表；同时分析了 createImportedRecord() 方法作为文档导入专用入口的特殊设计。
（2）ProfileInsightService 详细分析：分析了 getCurrentProfileDetail() 方法如何聚合个人档案、最新记录、6项趋势指标（buildTrends() 方法计算体重/收缩压/空腹血糖/睡眠/运动/风险分数的 latest/previous/change/direction/interpretation）、最多6条个性化建议（buildPersonalizedSuggestions() 基于趋势阈值和档案信息生成）以及近10条记录和预警；分析了 buildConsultationContext() 方法如何将档案信息、健康记录、趋势摘要、系统建议和用户问题组装为结构化的 AI 提示词上下文。
（3）DocumentImportService 详细分析：分析了文档导入的完整流程——extractText() 根据文件类型分发到PDF提取（PDFBox）或图片OCR提取（系统级 Tesseract 进程，参数 -l chi_sim+eng）或纯文本读取；parseDocument() 使用26个正则表达式提取字段（支持中英文字段名，如"血压/Blood Pressure"、"空腹血糖/Fasting Glucose"等）；mergeProfile() 采用"非空优先"策略将解析结果合并入现有档案。
（4）VisualizationService 分析：分析了 getVisualization() 如何取最近20条记录构建6个 VisualizationSeriesResponse 系列（每系列含 metricCode/metricName/unit/latest/average/points），以及如何同时提供风险等级分布Map和近10条记录摘要。
（5）异常处理机制：分析了 UnauthorizedException（401未授权）、NotFoundException（404资源不存在）、BadRequestException（400参数错误）、AiIntegrationException（AI服务异常）、ImportProcessingException（文档导入异常）5类自定义异常及其在各Service中的使用场景。
自我评价：
通过对关键服务类和工具类的深入阅读，对系统的核心业务逻辑有了完整把握。尤其是对健康记录保存、风险评估和预警刷新的完整业务链路，以及文档导入的多格式处理流程，形成了清晰的代码级认识，为后续撰写详细设计说明书中的业务流程描述部分提供了充分依据。"""

for cell in table5.rows[1].cells:
    if cell.text.strip().startswith('本阶段完成的具体工作'):
        set_cell_text(cell, new_t5r1_text)
        break


# ======================================================
# Table 6: 工作进程记录4（第10周-第13周）
# ======================================================
table6 = doc.tables[6]

new_t6r1_text = """本阶段完成的具体工作回顾及完成情况自我评价。
工作回顾：
本阶段对系统主要功能进行了全面联调与测试分析，验证了各模块在实际运行中的正确性与一致性。
主要测试验证内容如下：
（1）用户认证测试：验证了用户注册（/api/auth/register）时密码一致性校验、用户名唯一性校验和 BCrypt 加密存储；登录（/api/auth/login）时密码哈希比对；会话状态查询（/api/auth/me）在已登录/未登录状态下的正确响应；退出登录（/api/auth/logout）后 Session 失效的行为。
（2）健康记录管理测试：测试了记录的CRUD操作，验证了保存记录后 RiskAssessmentService 自动触发的风险评分计算结果，确认不同指标组合下风险等级（LOW/MEDIUM/HIGH/CRITICAL）的正确映射；验证了删除记录时级联删除关联预警的行为（healthAlertRepository.deleteByHealthRecordId()）；测试了按 riskLevel 筛选分页查询的正确性。
（3）健康预警测试：验证了预警在记录保存后的自动生成和刷新逻辑（refreshAlerts() 先删后增）；测试了预警状态更新（PENDING→其他状态）和 handledNote 处理备注的保存；验证了按 status/severity 组合筛选的4种查询场景。
（4）数据可视化测试：验证了 VisualizationService 取最近20条记录的逻辑，确认了6个指标系列在有/无数据情况下的正确返回，以及 latest/average 的计算准确性。
（5）文档导入测试：测试了纯文本格式（含"血压: 120/80"、"空腹血糖: 5.6"等标准字段）的解析正确性，验证了 mergeProfile() 非空优先合并策略；对PDF类型使用 PDFBox 提取文本后的正常解析流程进行了验证；确认了在无可识别字段时抛出 ImportProcessingException 的预期行为。
（6）档案详情与AI咨询测试：验证了 ProfileInsightService 对6项趋势指标方向判断（稳定/上升/下降）的正确性；验证了个性化建议生成逻辑（空腹血糖上升≥0.5时触发建议，收缩压上升≥8时触发建议等）；对 DeepSeekConsultationService 在 deepseek.enabled=false 时抛出 AiIntegrationException 的降级行为进行了验证。
（7）异常处理验证：验证了 GlobalExceptionHandler 对各类自定义异常的响应状态码映射和错误信息格式，确认未登录访问业务接口时返回401状态码的行为一致性。
自我评价：
本阶段能较好地把前后端交互与业务逻辑结合起来进行系统性验证，对系统可运行性、可演示性和代码完整性有了较明确的判断。认识到AI功能依赖外部DeepSeek API Key和网络环境，答辩时需要如实说明触发条件，并准备好模拟演示方案。整体测试过程中未发现主要功能性缺陷，系统状态稳定、可演示。"""

for cell in table6.rows[1].cells:
    if cell.text.strip().startswith('本阶段完成的具体工作'):
        set_cell_text(cell, new_t6r1_text)
        break


# ======================================================
# Table 7: 工作进程记录5（第14周-第16周）
# ======================================================
table7 = doc.tables[7]

new_t7r1_text = """本阶段完成的具体工作回顾及完成情况自我评价。
工作回顾：
本阶段主要完成毕业设计说明书和手册整理工作，以源码为唯一事实依据，对各部分内容进行系统性归纳与核实。
主要工作内容如下：
（1）需求分析部分撰写：基于对源码的完整梳理，撰写了系统功能需求（8大功能模块20个API端点）、非功能需求（安全性：BCrypt+Session认证；性能：readOnly事务优化；可靠性：统一异常处理；可扩展性：AI功能开关设计）和用例描述。
（2）总体设计部分撰写：绘制并描述了系统整体架构（前端层/Controller层/Service层/Repository层/数据库层）、模块依赖关系（DashboardService依赖AuthService/UserProfileService/HealthRecordRepository/HealthAlertRepository；HealthRecordService依赖AuthService/UserProfileService/RiskAssessmentService/HealthAlertRepository等）以及数据库4张表的E-R关系图。
（3）详细设计部分撰写：重点撰写了RiskAssessmentService的12类指标评估算法说明（含具体阈值表格：BMI≥24.0触发MEDIUM级别体重预警，≥28.0触发HIGH级别；收缩压≥130触发MEDIUM，≥140触发HIGH，≥180触发CRITICAL等）；撰写了DocumentImportService的文档解析流程说明（PDF提取→正则解析26字段→档案合并→记录创建）；撰写了DeepSeekConsultationService的AI集成说明（systemPrompt设计、context构建、temperature/max_tokens参数说明）。
（4）数据库设计部分撰写：梳理并记录了4张表（app_user/user_profile/health_record/health_alert）的完整字段定义、数据类型、约束条件和外键关系，与BaseEntity公共字段一并说明。
（5）系统测试部分撰写：整理了认证模块、记录管理模块、预警模块、可视化模块、文档导入模块和AI咨询模块的测试用例与测试结果。
（6）答辩材料准备：整理了系统演示顺序（用户注册→档案完善→记录录入→预警查看→可视化→详情分析→文档导入），准备了主要代码段说明（RiskAssessmentService评估逻辑、DeepSeekConsultationService集成逻辑、DocumentImportService文档解析逻辑），归纳了项目创新点和待改进方向。
自我评价：
能够在文档撰写阶段坚持以源码为依据，对系统功能、技术栈和模块分工进行客观描述，避免了夸大或虚构功能的问题。通过整个毕业设计过程，对软件工程"需求—设计—实现—测试—文档"的完整流程有了更系统的认识，也深刻体会到文档与代码保持一致的重要性。"""

for cell in table7.rows[1].cells:
    if cell.text.strip().startswith('本阶段完成的具体工作'):
        set_cell_text(cell, new_t7r1_text)
        break


# ======================================================
# Table 8: 中期检查报告
# ======================================================
table8 = doc.tables[8]

# 已完成的工作 [3,0]
new_t8r3_text = """已完成的工作
根据任务书要求，对项目源码进行了系统梳理和分析，已经完成课题调研、开题报告、系统需求分析与总体设计；完成了对用户认证、个人健康档案管理、健康记录管理、风险评估与预警、健康总览仪表盘、数据可视化、健康文档导入和AI咨询等8大功能模块的实现过程分析；深入研读了 RiskAssessmentService（多指标阈值评估算法）、ProfileInsightService（档案详情聚合与趋势分析）、DocumentImportService（PDF/OCR/文本解析）和 DeepSeekConsultationService（AI集成）等核心Service类的实现细节；对前后端交互流程（20个REST API端点）、实体关系（4张数据表的关联设计）、主要接口和关键服务逻辑进行了整理，并对系统主要功能进行了联调与测试验证。"""

for cell in table8.rows[3].cells:
    if cell.text.strip().startswith('已完成的工作'):
        set_cell_text(cell, new_t8r3_text)
        break

# 取得的阶段性成果 [4,0]
new_t8r4_text = """取得的阶段性成果
1. 明确了项目的真实技术栈，包括 Spring Boot 3.5.11、Spring MVC、Spring Data JPA（Hibernate）、Spring Security Crypto（BCrypt）、Apache PDFBox 3.0.5、Tesseract OCR（系统级）、H2 文件数据库（MySQL兼容）、HTML/CSS/JavaScript 前端和 DeepSeek Chat API 接口。
2. 梳理了 BaseEntity、AppUser、UserProfile、HealthRecord、HealthAlert 等核心实体及其1:1、1:N关联关系，以及19个DTO类和4个Repository接口的设计。
3. 明确了健康记录保存后自动触发风险评估（RiskAssessmentService.assess()）和预警刷新（HealthAlertRepository.deleteByHealthRecordId()后批量保存新预警）的关键业务链路。
4. 梳理了 RiskAssessmentService 中对12类生理与生活方式指标（BMI/体温/空腹血糖/餐后血糖/血氧/总胆固醇/腰围/血压/睡眠/压力/步数/情绪）建立的分级阈值评估规则，以及风险等级（LOW/MEDIUM/HIGH/CRITICAL）的映射逻辑。
5. 分析了 DocumentImportService 的文档导入流程（PDF→PDFBox提取/图片→Tesseract OCR/文本→直读 → 26字段正则解析 → 档案合并 → 健康记录创建）。
6. 完成了开题报告、5个阶段工作进程记录和中期检查所需的阶段性文档整理。"""

for cell in table8.rows[4].cells:
    if cell.text.strip().startswith('取得的阶段性成果'):
        set_cell_text(cell, new_t8r4_text)
        break

# 存在的问题 [5,0]
new_t8r5_text = """存在的问题
1. AI 咨询功能依赖外部 DeepSeek API Key（通过 DEEPSEEK_API_KEY 环境变量注入）和网络环境，演示时需根据配置情况说明其可用条件；在 deepseek.enabled=false（默认）时，调用会抛出 AiIntegrationException，需要提前说明。
2. 图片OCR功能依赖系统级 Tesseract OCR 安装及中文语言包（chi_sim），在未安装 Tesseract 的环境中图片导入功能不可用，需在演示环境中提前配置。
3. 系统采用 H2 文件数据库与演示数据初始化方式，更适合课程设计和毕业设计展示，若用于更大规模场景仍需替换为生产级数据库并引入连接池配置。
4. 风险评估当前采用规则阈值法（RuleBasedAssessment），评估覆盖12类指标但每类仅区分2-3个阈值区间，后续仍可继续完善评估指标细粒度与动态权重策略。"""

for cell in table8.rows[5].cells:
    if cell.text.strip().startswith('存在的问题'):
        set_cell_text(cell, new_t8r5_text)
        break

# 下一步工作计划 [6,0]
new_t8r6_text = """下一步工作计划
继续完善毕业设计说明书和手册内容，重点补充系统架构图（分层架构与模块依赖关系）、数据库设计（4张表的完整字段说明与E-R图）、关键业务流程说明（健康记录保存→风险评估→预警刷新的详细流程图、文档导入的多格式处理流程图）和测试结果；结合答辩需要整理各模块演示顺序（注册→档案→记录→预警→可视化→详情→文档导入→AI咨询）、关键代码说明（RiskAssessmentService评估逻辑、DocumentImportService解析逻辑）和项目总结，确保文档表述与实际源码保持一致。"""

for cell in table8.rows[6].cells:
    if cell.text.strip().startswith('下一步工作计划'):
        set_cell_text(cell, new_t8r6_text)
        break


# ======================================================
# Table 9: 毕业设计工作总结
# ======================================================
table9 = doc.tables[9]

# 工作任务完成情况 [0,0]
new_t9r0_text = """工作任务完成情况（包括任务书中规定的工作内容、研究目标等，如未完成应说明原因）
已按照任务书要求完成毕业设计主要工作。系统围绕个人健康管理场景，基于 Spring Boot 3.5.11 和 Java 21 实现了8大功能模块：用户注册登录（BCrypt加密+Session认证）、个人健康档案管理（24个字段的用户档案CRUD）、健康记录管理（22项生理与生活方式指标的录入/修改/删除/分页查询）、风险评估与预警（12类指标多阈值评估，自动生成 LOW/MEDIUM/HIGH/CRITICAL 四级预警）、健康总览仪表盘（聚合展示关键统计数据与最近记录/预警）、数据可视化（最近20条记录的6指标趋势系列+风险分布）、健康文档导入（PDF/OCR/文本三类格式，解析26个字段）以及可选AI咨询接口（集成 DeepSeek Chat API）。后端完成了分层结构设计与20个REST API端点实现，前端完成了页面展示、表单交互、筛选与分页逻辑；已对主要业务流程进行了测试与整理，并完成相关设计文档撰写。"""

for cell in table9.rows[0].cells:
    if cell.text.strip().startswith('工作任务完成情况'):
        set_cell_text(cell, new_t9r0_text)
        break

# 主要创新点 [1,0]
new_t9r1_text = """主要创新点：
1. 围绕单一用户建立了完整的健康管理闭环——用户注册→档案完善→持续健康记录→自动风险评估→预警生成→趋势分析→AI咨询，而不是简单的数据录入页面。
2. 在健康记录保存后自动触发多指标风险评估服务（RiskAssessmentService），对12类生理与生活方式指标进行分级阈值判断，依据综合评分和最高严重度映射风险等级，并自动刷新对应的健康预警列表，实现了记录与预警的联动管理。
3. 通过 ProfileInsightService 在档案详情接口中聚合最新健康快照、6项近期趋势指标（含方向判断与解读文本）、基于趋势阈值和档案特征的个性化健康建议以及近期预警，提升了系统信息的综合展示能力与个性化程度。
4. 实现了多格式健康文档导入功能（DocumentImportService），支持PDF文本提取（Apache PDFBox）、图片OCR识别（Tesseract，支持中英文混排）和纯文本三类输入，通过26个正则表达式自动解析健康字段并合并入用户档案，降低了健康数据录入门槛。
5. 集成 DeepSeek Chat API（通过 Spring RestClient 调用），以结构化提示词将个人档案、健康记录、趋势摘要和系统建议组装为上下文，实现了AI辅助健康分析功能，并通过环境变量开关（DEEPSEEK_ENABLED）支持灵活配置。"""

for cell in table9.rows[1].cells:
    if cell.text.strip().startswith('主要创新点'):
        set_cell_text(cell, new_t9r1_text)
        break

# 工作状况 [2,0]
new_t9r2_text = """工作状况（包括工作态度、刻苦精神、协作精神、个人精力投入、出勤情况等）：
毕业设计期间能够按照计划推进各阶段任务，重视资料查阅、源码分析、系统测试与文档整理等工作。对项目中的核心模块（特别是 RiskAssessmentService、ProfileInsightService、DocumentImportService 和 DeepSeekConsultationService）进行了细致的阅读与理解，能够在发现文档描述与代码实现不一致时及时修正，保证文档内容与实际代码保持一致。整体工作态度认真，各阶段任务均按时推进，投入较为稳定。"""

for cell in table9.rows[2].cells:
    if cell.text.strip().startswith('工作状况'):
        set_cell_text(cell, new_t9r2_text)
        break

# 收获、体会及建议 [3,0]
new_t9r3_text = """收获、体会及建议：
通过本次毕业设计，进一步加深了对 Spring Boot Web 应用开发完整流程的理解，掌握了实体设计（JPA实体关系映射）、接口定义（RESTful API规范）、业务服务封装（事务管理与只读优化）、认证与安全（BCrypt+Session）、前后端交互（JSON序列化/反序列化与跨域配置）、文档处理（PDFBox+Tesseract集成）、外部API调用（RestClient+DeepSeek）和异常处理（全局异常处理器）等实践能力；也深刻认识到在毕业设计写作中必须坚持以真实代码为依据，不能脱离项目实际随意扩展功能描述，否则在答辩时难以自圆其说。后续若继续完善本系统，可在以下方面继续拓展：基于 Spring Security 引入更完善的权限控制与Token认证；集成 ECharts 图表库改善数据可视化展示效果；引入 Spring Boot Actuator 增加系统监控能力；在更大规模场景下替换 H2 为 PostgreSQL 并引入 Flyway 数据库版本管理；进一步细化风险评估模型，引入动态权重或机器学习模型提升评估准确性。"""

for cell in table9.rows[3].cells:
    if cell.text.strip().startswith('收获、体会及建议'):
        set_cell_text(cell, new_t9r3_text)
        break


# ======================================================
# 保存文档
# ======================================================
output_filename = '基于SpringBoot的个人健康管理与疾病预警系统（更新版）.docx'
doc.save(output_filename)
print(f"文档已保存为: {output_filename}")
print("主要修改内容：")
print("1. 任务书（Table 0）：更新功能模块描述（家庭成员→个人用户，补充认证/可视化/文档导入模块）")
print("2. 任务书（Table 0）：更新技术指标（补充PDFBox、Tesseract、BCrypt、AI参数等）")
print("3. 开题报告（Table 2）：更新课题目的意义、资料调研分析、可行性分析")
print("4. 工作进程记录1（Table 3）：扩充第1-3周详细工作内容")
print("5. 工作进程记录2（Table 4）：扩充第4-6周详细工作内容")
print("6. 工作进程记录3（Table 5）：扩充第7-9周详细工作内容")
print("7. 工作进程记录4（Table 6）：扩充第10-13周详细工作内容（含各模块测试验证）")
print("8. 工作进程记录5（Table 7）：扩充第14-16周详细工作内容（含说明书各章节撰写）")
print("9. 中期检查报告（Table 8）：更新已完成工作、阶段性成果、存在问题、下一步计划")
print("10. 毕业设计工作总结（Table 9）：更新任务完成情况、创新点、工作状况、收获体会")
